from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import authenticate, login as auth_login, logout as auth_logout
from django.contrib.auth.decorators import login_required
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse, HttpResponseNotFound
from django.http import Http404
from .models import Teacher, Question, Subject
from .forms import QuestionForm
from .utils import (
    extract_text_from_pdf,
    clean_text,
    generate_questions,
    recognize_and_generate_diagram_questions,
    save_questions_to_word,
)
import os
import random
from io import BytesIO
from docx import Document

def home(request):
    if request.user.is_authenticated:  # Check if the user is already logged in
        return redirect('dashboard')
    
    return render(request, 'main/home.html')

from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Teacher

def register(request):
    if request.user.is_authenticated:  
        return redirect('dashboard')    
    
    if request.method == 'POST':
        first_name = request.POST['first_name']
        last_name = request.POST['last_name']
        email = request.POST['email']
        specialization = request.POST['specialization']
        password = request.POST['password']
        confirm_password = request.POST['confirm_password']
        profile_photo = request.FILES.get('profile_photo')  # Get uploaded file

        if password != confirm_password:
            messages.error(request, "Passwords do not match. Please try again.")
            return redirect('register')

        if Teacher.objects.filter(email=email).exists():
            messages.error(request, "Email already exists. Please use a different one.")
            return redirect('register')

        new_teacher = Teacher.objects.create_user(
            email=email, 
            password=password, 
            first_name=first_name,
            last_name=last_name,
            specialization=specialization
        )

        if profile_photo:
            new_teacher.profile_photo = profile_photo
            new_teacher.save()

        messages.success(request, "Registration successful. Please log in.")
        return redirect('login')
    
    return render(request, 'main/register.html')

def login(request):
    if request.user.is_authenticated:  # Check if the user is already logged in
        return redirect('dashboard')

    if request.method == 'POST':
        email = request.POST['email']
        password = request.POST['password']
        user = authenticate(request, email=email, password=password)

        if user is not None:
            auth_login(request, user)
            messages.success(request, "Login successful!")
            return redirect('dashboard')
        else:
            messages.error(request, "Invalid email or password.")
            return redirect('login')

    return render(request, 'main/login.html')

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from main.models import Teacher

@login_required
def profile(request):
    teacher = request.user  # Get the logged-in teacher

    if request.method == "POST":
        # Handle profile updates
        teacher.first_name = request.POST.get("first_name", teacher.first_name)
        teacher.last_name = request.POST.get("last_name", teacher.last_name)
        teacher.specialization = request.POST.get("specialization", teacher.specialization)
        
        if "profile_photo" in request.FILES:
            teacher.profile_photo = request.FILES["profile_photo"]
        
        teacher.save()
        messages.success(request, "Profile updated successfully!")

        return redirect("profile")  # Refresh the page after updating

    return render(request, "main/profile.html", {"teacher": teacher})


def logout(request):
    auth_logout(request)
    messages.success(request, "Logged out successfully.")
    return redirect('login')

def dashboard(request):
    if request.user.is_authenticated:
        return render(request, 'main/dashboard.html')
    else:
        messages.error(request, "You need to log in first.")
        return redirect('login')
    

@login_required
def insert_questions(request):
    if request.method == 'POST':
        semester = request.POST.get('semester')
        subject_name = request.POST.get('subject_name')
        subject_code = request.POST.get('subject_code')

        # Get the logged-in teacher
        teacher_email = request.user.email
        teacher = get_object_or_404(Teacher, email=teacher_email)

        # Check if the subject already exists for the logged-in teacher
        existing_subject = Subject.objects.filter(
            teacher=teacher, subject_name=subject_name, subject_code=subject_code
        ).first()

        if existing_subject:
            messages.success(request, "Subject already exists. You can now add questions.")
            return redirect('add_question', subject_id=existing_subject.id)

        # Create a new subject
        new_subject = Subject(
            semester=semester,
            subject_name=subject_name,
            subject_code=subject_code,
            teacher=teacher
        )

        try:
            new_subject.save()
            messages.success(request, "Subject added successfully. Now you can add questions.")
            return redirect('add_question', subject_id=new_subject.id)
        except Exception as e:
            messages.error(request, f"An error occurred: {str(e)}")
            return redirect('insert_questions')

    return render(request, 'main/insert_questions.html')

@login_required
def generate_paper_form(request):
    return render(request, 'main/generate_paper_form.html')

@login_required
def add_question(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    
    if request.method == 'POST':
        form = QuestionForm(request.POST)
        if form.is_valid():
            question = form.save(commit=False)
            question.subject = subject
            question.save()
            messages.success(request, "Question added successfully!")
            # return redirect('subject_detail', subject_id=subject.id)  # Adjust this URL as needed
        else:
            messages.error(request, "An error occurred while adding the question.")
    else:
        form = QuestionForm()

    return render(request, 'main/add_question.html', {'form': form, 'subject': subject})


@login_required
def view_questions(request, subject_id):
    # Ensure the teacher is logged in
    if not request.user.is_authenticated:
        messages.warning(request, "Please log in to view the questions.")
        return redirect('login')

    # Query for all questions linked to the current subject
    questions = Question.objects.filter(subject_id=subject_id)
    
    # Get the subject object
    subject = get_object_or_404(Subject, id=subject_id)
    
    # Get unique chapter numbers in sorted order
    chapters = sorted(questions.values_list('chapter_no', flat=True).distinct())
    
    # Create dictionaries to track if chapters have questions of each mark value
    chapter_2_marks_exists = {chapter: False for chapter in chapters}
    chapter_5_marks_exists = {chapter: False for chapter in chapters}
    chapter_10_marks_exists = {chapter: False for chapter in chapters}
    
    # Check which chapters have which type of questions
    for question in questions:
        if question.marks == 2:
            chapter_2_marks_exists[question.chapter_no] = True
        elif question.marks == 5:
            chapter_5_marks_exists[question.chapter_no] = True
        elif question.marks == 10:
            chapter_10_marks_exists[question.chapter_no] = True
    
    context = {
        'questions': questions,
        'subject': subject,
        'chapters': chapters,
        'chapter_2_marks_exists': chapter_2_marks_exists,
        'chapter_5_marks_exists': chapter_5_marks_exists,
        'chapter_10_marks_exists': chapter_10_marks_exists,
    }
    
    return render(request, 'main/view_questions.html', context)


@login_required
def edit_question(request, question_id):
    # Fetch the question by ID or return a 404 if not found
    question = get_object_or_404(Question, id=question_id)

    # Check if the form is being submitted (POST request)
    if request.method == 'POST':
        form = QuestionForm(request.POST, instance=question)
        
        # Check if the form is valid
        if form.is_valid():
            try:
                form.save()  # Save the updated question
                messages.success(request, "Question updated successfully!")
                return redirect('view_questions', subject_id=question.subject.id)  # Redirect to the subject's questions page
            except Exception as e:
                messages.error(request, f"Error! Unable to update the question. {e}")
        else:
            messages.error(request, "Invalid form submission.")

    else:
        form = QuestionForm(instance=question)  # Pre-populate the form with the existing question details

    return render(request, 'main/edit_question.html', {'form': form, 'question': question})


@login_required
def update_question(request, question_id):
    # Get the question to update
    question = get_object_or_404(Question, id=question_id)
    
    if request.method == 'POST':
        form = QuestionForm(request.POST, instance=question)
        
        if form.is_valid():
            try:
                form.save()  # Save the updated question
                messages.success(request, "Question updated successfully!")
                return redirect('view_questions', subject_id=question.subject.id)  # Redirect to the subject's questions page
            except Exception as e:
                messages.error(request, f"Error! Unable to update the question. {e}")
        else:
            messages.error(request, "Invalid form submission.")
    
    else:
        form = QuestionForm(instance=question)  # Pre-populate the form with the existing question details

    return render(request, 'main/edit_question.html', {'form': form, 'question': question})

@login_required
def delete_question(request, question_id, subject_id):
    try:
        # Fetch the question by ID
        question_to_delete = get_object_or_404(Question, id=question_id)

        # Delete the question
        question_to_delete.delete()

        # Flash success message
        messages.success(request, "Question deleted successfully.")
        
    except Exception as e:
        # If an error occurs, send a failure message
        messages.error(request, f"An error occurred: {str(e)}")
    
    # Redirect to the view questions page for the specified subject
    return redirect('view_questions', subject_id=subject_id)

def fallback_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')  # Redirect to dashboard if user is logged in
    else:
        return redirect('home')  # Redirect to home if user is not logged in
    
    
    
from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from io import BytesIO

def quick_notes(request):
    """
    View function for the Quick Notes feature.
    This renders the template with the client-side implementation
    of the notes functionality using localStorage.
    """
    return render(request, 'main/quick_notes.html')

def download_note(request):
    """
    View function to download a note as a Word document.
    Receives the note content and filename via POST and returns
    a Word document for download.
    """
    if request.method == 'POST':
        file_name = request.POST.get('file_name', 'note')
        content = request.POST.get('content', '')
        
        # Create a new Word document
        doc = Document()
        
        # Add a title with the file name
        doc.add_heading(file_name, level=1)
        
        # Add content paragraphs
        for paragraph in content.split('\n'):
            if paragraph.strip():  # Skip empty lines
                doc.add_paragraph(paragraph)
        
        # Save the document to a BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create the HTTP response with the document
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{file_name}.docx"'
        
        return response
    
    # Redirect back to the notes page if not a POST request
    return redirect('quick_notes')